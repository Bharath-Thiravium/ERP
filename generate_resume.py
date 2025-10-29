from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Header
name = doc.add_paragraph("Ilayaraja Ramasamy")
name.runs[0].bold = True
name.runs[0].font.size = Pt(16)
name.alignment = WD_ALIGN_PARAGRAPH.CENTER

title = doc.add_paragraph("Senior Full-Stack Developer | Enterprise Software Architect | End-to-End System Builder")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

contact = doc.add_paragraph(
    "📧 ilaiarajacse@gmail.com | 📞 +91-9080849708 | 🌐 LinkedIn: linkedin.com/in/ilayaraja-ramasamy | GitHub: github.com/ilayaraja"
)
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

# Professional Summary
doc.add_heading("Professional Summary", level=1)
doc.add_paragraph(
    "Innovative Full-Stack Software Developer with 10+ years of experience building high-performance enterprise applications. "
    "Successfully architected and developed a comprehensive SAP web application system single-handedly, demonstrating exceptional "
    "technical leadership and problem-solving capabilities. Expert in modern web technologies, AI integration, and Indian statutory "
    "compliance systems with proven track record of delivering scalable, secure, and maintainable solutions."
)

# Core Competencies
doc.add_heading("Core Competencies", level=1)
competencies = [
    "• Full-Stack Development: React, Django, TypeScript, Python, PostgreSQL",
    "• Enterprise Architecture: Multi-tenant SaaS, Microservices, RESTful APIs",
    "• AI/ML Integration: Face Recognition, Predictive Analytics, OpenCV",
    "• Mobile Development: React Native, Cross-platform Applications",
    "• DevOps & Cloud: Docker, CI/CD, GitHub Actions, Linux Administration",
    "• Database Design: PostgreSQL, Redis, Query Optimization, Indexing",
    "• Security: JWT Authentication, Encryption, OWASP Compliance",
    "• Indian Compliance: GST, PF, ESI, TDS, Labor Law Automation"
]
for comp in competencies:
    doc.add_paragraph(comp)

# Major Project
doc.add_heading("Major Project Achievement", level=1)

project_title = doc.add_paragraph("Enterprise SAP Web Application System")
project_title.runs[0].bold = True
project_title.runs[0].font.size = Pt(14)

doc.add_paragraph("Solo Developer & System Architect | 18 Months Development | ₹3-6 Crores Market Value")

doc.add_heading("Project Overview", level=2)
doc.add_paragraph(
    "Designed and developed a comprehensive enterprise-grade SAP system from scratch as a single developer, "
    "delivering a complete business management solution with advanced features including AI-powered analytics, "
    "multi-tenant architecture, and Indian statutory compliance."
)

# Technical Stack
doc.add_heading("Technical Architecture & Implementation", level=2)
tech_stack = [
    "• Backend: Django 5.2.6, Django REST Framework, PostgreSQL, Redis, Celery",
    "• Frontend: React 19.1.1, TypeScript, Vite, Ant Design, TanStack Query",
    "• Mobile: React Native 0.81.4, Face Recognition, GPS Tracking",
    "• DevOps: Docker, GitHub Actions, Nginx, Gunicorn, Linux",
    "• AI/ML: OpenCV, Face Recognition, Predictive Analytics",
    "• Security: JWT, 2FA, Encryption, Rate Limiting, Audit Logging"
]
for tech in tech_stack:
    doc.add_paragraph(tech)

# Core Modules
doc.add_heading("Business Modules Developed", level=2)
modules = [
    "• Finance Management: Complete invoicing, GST compliance, payment tracking, financial reporting",
    "• Human Resources: HRMS, payroll processing, attendance system, performance management",
    "• Inventory Management: Stock tracking, purchase management, warehouse operations",
    "• CRM System: Lead management, sales pipeline, customer analytics, marketing automation",
    "• Analytics Engine: Real-time dashboards, predictive analytics, custom reporting"
]
for module in modules:
    doc.add_paragraph(module)

# Key Features
doc.add_heading("Advanced Features & Innovations", level=2)
features = [
    "• AI-Powered Face Recognition: 99.5% accuracy attendance system with OpenCV integration",
    "• Multi-tenant Architecture: Complete data isolation with service-based access control",
    "• Indian Statutory Compliance: Automated GST, PF, ESI, TDS calculations and reporting",
    "• Real-time Analytics: Live dashboards with predictive business intelligence",
    "• Mobile Integration: React Native app with biometric attendance and GPS tracking",
    "• Security Framework: Multi-factor authentication, encryption, comprehensive audit trails"
]
for feature in features:
    doc.add_paragraph(feature)

# Technical Achievements
doc.add_heading("Technical Achievements", level=2)
achievements = [
    "• Database Design: 150+ optimized tables with strategic indexing and relationships",
    "• API Development: 200+ RESTful endpoints with comprehensive documentation",
    "• Performance: Sub-200ms response times supporting 1000+ concurrent users",
    "• Code Quality: 95%+ test coverage with zero critical security vulnerabilities",
    "• Scalability: Horizontal scaling architecture for enterprise-level growth"
]
for achievement in achievements:
    doc.add_paragraph(achievement)

# Business Impact
doc.add_heading("Business Impact & Value", level=2)
impact = [
    "• Cost Reduction: 70% reduction in manual administrative tasks",
    "• Compliance Automation: 100% automated statutory compliance reporting",
    "• Process Efficiency: 60% faster invoice processing and approval workflows",
    "• Revenue Potential: ₹2.1+ crores annual recurring revenue at scale",
    "• Market Value: ₹3-6 crores current valuation, ₹25-50 crores at scale"
]
for imp in impact:
    doc.add_paragraph(imp)

# Technical Skills
doc.add_heading("Technical Skills", level=1)
skills_categories = [
    "Programming Languages: Python, JavaScript, TypeScript, SQL, HTML5, CSS3",
    "Frameworks & Libraries: Django, React, React Native, Node.js, Express.js",
    "Databases: PostgreSQL, MySQL, Redis, MongoDB",
    "Cloud & DevOps: Docker, GitHub Actions, Nginx, Linux, AWS basics",
    "Tools & Technologies: Git, Postman, VS Code, Figma, Jira"
]
for skill in skills_categories:
    doc.add_paragraph(f"• {skill}")

# Education
doc.add_heading("Education", level=1)
doc.add_paragraph("• Bachelor of Engineering in Computer Science")
doc.add_paragraph("• Relevant Certifications in Full-Stack Development and Cloud Technologies")

# Professional Highlights
doc.add_heading("Professional Highlights", level=1)
highlights = [
    "• Solo architect and developer of enterprise SAP system worth ₹3-6 crores",
    "• Expert in Indian business compliance and statutory requirements",
    "• Proven ability to deliver complex projects independently with high quality",
    "• Strong problem-solving skills with focus on scalable, maintainable solutions",
    "• Experience with complete software development lifecycle from design to deployment"
]
for highlight in highlights:
    doc.add_paragraph(highlight)

# Save Word file
doc.save("Ilayaraja_Ramasamy_Resume.docx")
print("Resume saved as Ilayaraja_Ramasamy_Resume.docx")