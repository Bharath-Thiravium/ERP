import PyPDF2
import pandas as pd
import docx
import re
from typing import Dict, List, Any
import logging

logger = logging.getLogger(__name__)

class TemplateParser:
    """Parse uploaded templates to extract form structure and fields"""
    
    @staticmethod
    def parse_template(file_path: str, file_type: str) -> Dict[str, Any]:
        """Parse template file and extract form structure"""
        try:
            if file_type == 'pdf':
                return TemplateParser._parse_pdf(file_path)
            elif file_type == 'excel':
                return TemplateParser._parse_excel(file_path)
            elif file_type == 'word':
                return TemplateParser._parse_word(file_path)
            else:
                return TemplateParser._get_default_structure()
        except Exception as e:
            logger.error(f"Error parsing template: {str(e)}")
            return TemplateParser._get_default_structure()
    
    @staticmethod
    def _parse_pdf(file_path: str) -> Dict[str, Any]:
        """Extract fields from PDF"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
            
            # Extract table headers and structure
            fields = TemplateParser._extract_fields_from_text(text)
            return {
                'form_type': 'dynamic_pdf',
                'fields': fields,
                'structure': 'table',
                'source': 'pdf'
            }
        except Exception as e:
            logger.error(f"PDF parsing error: {str(e)}")
            return TemplateParser._get_default_structure()
    
    @staticmethod
    def _parse_excel(file_path: str) -> Dict[str, Any]:
        """Extract fields from Excel"""
        try:
            df = pd.read_excel(file_path, nrows=5)  # Read first 5 rows to get headers
            fields = []
            
            # Get column headers
            for col in df.columns:
                if pd.notna(col) and str(col).strip():
                    fields.append({
                        'name': str(col).strip(),
                        'type': 'text',
                        'required': False
                    })
            
            return {
                'form_type': 'dynamic_excel',
                'fields': fields,
                'structure': 'table',
                'source': 'excel'
            }
        except Exception as e:
            logger.error(f"Excel parsing error: {str(e)}")
            return TemplateParser._get_default_structure()
    
    @staticmethod
    def _parse_word(file_path: str) -> Dict[str, Any]:
        """Extract fields from Word document"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract table structure if exists
            fields = []
            for table in doc.tables:
                if table.rows:
                    header_row = table.rows[0]
                    for cell in header_row.cells:
                        if cell.text.strip():
                            fields.append({
                                'name': cell.text.strip(),
                                'type': 'text',
                                'required': False
                            })
                    break  # Use first table
            
            if not fields:
                # Extract from text if no tables
                fields = TemplateParser._extract_fields_from_text(text)
            
            return {
                'form_type': 'dynamic_word',
                'fields': fields,
                'structure': 'table',
                'source': 'word'
            }
        except Exception as e:
            logger.error(f"Word parsing error: {str(e)}")
            return TemplateParser._get_default_structure()
    
    @staticmethod
    def _extract_fields_from_text(text: str) -> List[Dict[str, Any]]:
        """Extract field names from text content"""
        # Common field patterns - enhanced for advance register
        field_patterns = [
            r'S\.?\s*No\.?',
            r'Employee\s+(?:ID|Name|Code)',
            r'Name\s*(?:&\s*Surname)?',
            r'Date\s+of\s+(?:Birth|Joining|Commencement|Termination)',
            r'Father\'?s?\s*/?Husband\'?s?\s*Name',
            r'Nature\s+of\s+Employment',
            r'Permanent\s+Address',
            r'Local\s+Address',
            r'Department',
            r'Designation',
            r'Basic\s+(?:Wage|Salary)',
            r'Fine\s+Amount',
            r'Advance\s+Amount',
            r'Purpose',
            r'Installments?',
            r'Monthly\s+Deduction',
            r'Balance\s+Outstanding',
            r'Sex|Gender',
            r'Remarks?',
            # Additional patterns for advance register
            r'Amount\s+of\s+Advance',
            r'Date\s+of\s+Advance',
            r'Purpose\s+of\s+Advance',
            r'No\.?\s+of\s+Installments?',
            r'Amount\s+of\s+(?:each\s+)?Installment',
            r'Date\s+of\s+Recovery',
            r'Amount\s+Recovered',
            r'Balance\s+(?:Amount\s+)?Outstanding',
            r'Signature\s+of\s+Employee',
            r'Date\s+of\s+Repayment',
            r'Recovery\s+Details'
        ]
        
        fields = []
        for pattern in field_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                field_name = match.strip()
                if field_name and field_name not in [f['name'] for f in fields]:
                    fields.append({
                        'name': field_name,
                        'type': 'text',
                        'required': False
                    })
        
        # Add default fields if none found
        if not fields:
            fields = TemplateParser._get_default_fields()
        
        return fields
    
    @staticmethod
    def _get_default_fields() -> List[Dict[str, Any]]:
        """Default fields when parsing fails"""
        return [
            {'name': 'S.No', 'type': 'number', 'required': True},
            {'name': 'Employee ID', 'type': 'text', 'required': True},
            {'name': 'Employee Name', 'type': 'text', 'required': True},
            {'name': 'Department', 'type': 'text', 'required': False},
            {'name': 'Designation', 'type': 'text', 'required': False},
            {'name': 'Date of Joining', 'type': 'date', 'required': False},
            {'name': 'Remarks', 'type': 'text', 'required': False}
        ]
    
    @staticmethod
    def _get_default_structure() -> Dict[str, Any]:
        """Default structure when parsing fails"""
        return {
            'form_type': 'dynamic_template',
            'fields': TemplateParser._get_default_fields(),
            'structure': 'table',
            'source': 'default'
        }